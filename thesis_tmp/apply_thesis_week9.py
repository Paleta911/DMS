# -*- coding: utf-8 -*-
from pathlib import Path
import shutil
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

sys.stdout.reconfigure(encoding='utf-8')

BASE = Path(r'C:\Users\alexi\Documents\Estadías\estadias\PARA _CODEX')
TMP = Path(r'c:\Users\alexi\DMS\thesis_tmp')
DOC_PATH = BASE / 'Estadías v2.9.docx'
BACKUP_PATH = BASE / 'Estadías v2.9 - respaldo antes semana 9.docx'
ARCH_IMG = BASE / '_anexo_arquitectura_semana9.png'
ERD_IMG = BASE / '_anexo_erd_semana9.png'

PROJECT_TITLE = 'Sistema digital de gestión documental para el área de Calidad de Central El Potrero'
COVER_DATE = 'CUITLÁHUAC, VER.                                                                            MARZO, 2026'


def create_architecture_image(path: Path):
    w, h = 1800, 900
    img = Image.new('RGB', (w, h), '#f7f9fc')
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.truetype(r'C:\Windows\Fonts\arialbd.ttf', 42)
    label_font = ImageFont.truetype(r'C:\Windows\Fonts\arialbd.ttf', 30)
    text_font = ImageFont.truetype(r'C:\Windows\Fonts\arial.ttf', 24)
    draw.text((w // 2, 60), 'Arquitectura general del sistema DMS SIG', anchor='mm', font=title_font, fill='#183153')
    boxes = [
        ((110, 320, 420, 500), '#dbeafe', 'Usuario', 'Consulta, registra y da seguimiento'),
        ((520, 300, 920, 520), '#e0f2fe', 'Frontend', 'React + Vite\nInterfaz, formularios y navegación'),
        ((1030, 270, 1430, 550), '#dcfce7', 'Backend', 'NestJS\nReglas de negocio, seguridad\ny orquestación'),
        ((1520, 180, 1710, 360), '#fee2e2', 'SQL Server', 'Persistencia transaccional'),
        ((1520, 470, 1710, 650), '#fef3c7', 'Elasticsearch', 'Búsqueda avanzada e indexación'),
    ]
    for (x1, y1, x2, y2), color, title, subtitle in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=28, fill=color, outline='#94a3b8', width=3)
        draw.text(((x1 + x2) // 2, y1 + 55), title, anchor='mm', font=label_font, fill='#0f172a')
        draw.multiline_text(((x1 + x2) // 2, (y1 + y2) // 2 + 15), subtitle, anchor='mm', font=text_font, fill='#334155', align='center', spacing=8)
    def arrow(points, label):
        draw.line(points, fill='#2563eb', width=8)
        x1, y1, x2, y2 = points
        draw.polygon([(x2, y2), (x2 - 20, y2 - 12), (x2 - 20, y2 + 12)], fill='#2563eb')
        draw.rounded_rectangle(((x1 + x2) // 2 - 95, (y1 + y2) // 2 - 30, (x1 + x2) // 2 + 95, (y1 + y2) // 2 + 30), radius=12, fill='#ffffff', outline='#cbd5e1')
        draw.text(((x1 + x2) // 2, (y1 + y2) // 2), label, anchor='mm', font=text_font, fill='#1e3a8a')
    arrow((420, 410, 520, 410), 'Uso web')
    arrow((920, 410, 1030, 410), 'HTTP / JSON')
    arrow((1430, 320, 1520, 270), 'ORM / SQL')
    arrow((1430, 500, 1520, 560), 'Índice / búsqueda')
    draw.text((w // 2, 820), 'El backend concentra la seguridad, la lógica de negocio y la comunicación con la base de datos y el motor de búsqueda.', anchor='mm', font=text_font, fill='#475569')
    img.save(path)


def create_erd_image(path: Path):
    w, h = 2100, 1300
    img = Image.new('RGB', (w, h), '#f8fafc')
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.truetype(r'C:\Windows\Fonts\arialbd.ttf', 40)
    head_font = ImageFont.truetype(r'C:\Windows\Fonts\arialbd.ttf', 24)
    text_font = ImageFont.truetype(r'C:\Windows\Fonts\arial.ttf', 20)
    small_font = ImageFont.truetype(r'C:\Windows\Fonts\arial.ttf', 18)
    draw.text((w // 2, 50), 'Modelo entidad-relación general del sistema', anchor='mm', font=title_font, fill='#1e293b')
    boxes = {
        'user': ((140, 180, 430, 340), '#ede9fe', ['id (PK)', 'email', 'role', 'status']),
        'email_verification': ((560, 180, 930, 340), '#e0f2fe', ['id (PK)', 'userId (FK)', 'codeHash', 'expiresAt']),
        'permission_request': ((1060, 180, 1450, 360), '#fee2e2', ['id (PK)', 'userId (FK)', 'reviewedById (FK)', 'status']),
        'audit_log': ((1550, 180, 1900, 340), '#dcfce7', ['id (PK)', 'userId (FK)', 'action', 'createdAt']),
        'category': ((140, 520, 430, 650), '#fde68a', ['id (PK)', 'nombre', 'isArchived']),
        'document_type': ((540, 500, 900, 670), '#fef3c7', ['id (PK)', 'code', 'nombreLargo', 'isArchived']),
        'area_code': ((1020, 500, 1370, 670), '#dbeafe', ['id (PK)', 'code', 'nombre', 'isArchived']),
        'user_area_codes': ((1470, 500, 1870, 650), '#e9d5ff', ['userId (FK)', 'areaCodeId (FK)']),
        'document': ((520, 860, 900, 1080), '#ccfbf1', ['id (PK)', 'categoryId (FK)', 'documentTypeId (FK)', 'areaCodeId (FK)', 'createdById (FK)']),
        'version': ((1000, 900, 1320, 1080), '#fee2e2', ['id (PK)', 'documentId (FK)', 'uploadedById (FK)', 'createdAt']),
        'document_approval': ((1450, 860, 1900, 1080), '#dbeafe', ['id (PK)', 'documentId (FK)', 'userId (FK)', 'step', 'decision']),
    }
    def box(name, rect, color, fields):
        x1, y1, x2, y2 = rect
        draw.rounded_rectangle(rect, radius=22, fill=color, outline='#94a3b8', width=3)
        draw.rectangle((x1, y1, x2, y1 + 42), fill='#1e293b')
        draw.text((x1 + 18, y1 + 21), name, anchor='lm', font=head_font, fill='white')
        y = y1 + 68
        for field in fields:
            draw.text((x1 + 18, y), field, anchor='lm', font=text_font, fill='#334155')
            y += 32
    for name, (rect, color, fields) in boxes.items():
        box(name, rect, color, fields)
    def connect(a_name, b_name, label=''):
        ax1, ay1, ax2, ay2 = boxes[a_name][0]
        bx1, by1, bx2, by2 = boxes[b_name][0]
        start = ((ax1 + ax2) // 2, ay2)
        end = ((bx1 + bx2) // 2, by1)
        draw.line((start[0], start[1], start[0], start[1] + 40, end[0], end[1] - 40, end[0], end[1]), fill='#2563eb', width=5)
        if label:
            lx = (start[0] + end[0]) // 2
            ly = (start[1] + end[1]) // 2
            draw.rounded_rectangle((lx - 45, ly - 18, lx + 45, ly + 18), radius=9, fill='white', outline='#cbd5e1')
            draw.text((lx, ly), label, anchor='mm', font=small_font, fill='#1e40af')
    connect('user', 'email_verification', '1:1')
    connect('user', 'permission_request', '1:N')
    connect('user', 'audit_log', '1:N')
    connect('category', 'document', '1:N')
    connect('document_type', 'document', '1:N')
    connect('area_code', 'document', '1:N')
    connect('area_code', 'user_area_codes', '1:N')
    connect('user', 'user_area_codes', '1:N')
    connect('document', 'version', '1:N')
    connect('document', 'document_approval', '1:N')
    draw.text((w // 2, 1230), 'El modelo organiza documentos, versiones, aprobaciones, auditoría, usuarios, áreas y solicitudes de permisos dentro de un esquema relacional centralizado.', anchor='mm', font=small_font, fill='#475569')
    img.save(path)


def load_items(name):
    items = []
    for line in (TMP / name).read_text(encoding='utf-8').splitlines():
        line = line.strip()
        if not line:
            continue
        kind, text = line.split('|', 1)
        items.append((kind, text))
    return items


def find_paragraph(doc, needle, exact=True):
    for p in doc.paragraphs:
        text = p.text.strip()
        if (text == needle) if exact else (needle in text):
            return p
    raise RuntimeError(f'No se encontró: {needle}')


def remove_between(start_paragraph, end_paragraph, include_start=False):
    body = start_paragraph._element.getparent()
    children = list(body)
    start_idx = children.index(start_paragraph._element) + (0 if include_start else 1)
    end_idx = children.index(end_paragraph._element)
    for child in children[start_idx:end_idx]:
        body.remove(child)


def add_item_before(anchor, kind, text):
    p = anchor.insert_paragraph_before('')
    if kind in {'H2', 'H3'}:
        p.text = text
        p.style = 'Heading 2' if kind == 'H2' else 'Heading 3'
    elif kind == 'P':
        p.text = text
    elif kind == 'FIG':
        run = p.add_run()
        run.add_picture(str(ARCH_IMG if text == 'ARCH' else ERD_IMG), width=Cm(14.5 if text == 'ARCH' else 15.5))
    elif kind == 'CAP':
        p.text = text
    elif kind == 'SRC':
        p.text = text
    return p




def find_next_sdt_after(paragraph):
    body = paragraph._element.getparent()
    children = list(body)
    idx = children.index(paragraph._element) + 1
    for child in children[idx:]:
        if child.tag.endswith('}sdt'):
            return child
    raise RuntimeError('No se encontró el bloque de bibliografía')


def insert_items_after(anchor_paragraph, items):
    current = anchor_paragraph
    inserted = []
    for kind, text in items:
        new_p = current.insert_paragraph_before('')
        current._element.addnext(new_p._element)
        current = new_p
        if kind in {'H2', 'H3'}:
            new_p.text = text
            new_p.style = 'Heading 2' if kind == 'H2' else 'Heading 3'
        elif kind == 'P':
            new_p.text = text
        elif kind == 'FIG':
            run = new_p.add_run()
            run.add_picture(str(ARCH_IMG if text == 'ARCH' else ERD_IMG), width=Cm(14.5 if text == 'ARCH' else 15.5))
        elif kind == 'CAP':
            new_p.text = text
        elif kind == 'SRC':
            new_p.text = text
        inserted.append((new_p, kind))
    return inserted

def fmt(paragraph, kind='P'):
    pf = paragraph.paragraph_format
    if kind == 'H1':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size, bold, italic = 14, True, False
        pf.space_before, pf.space_after = Pt(12), Pt(12)
    elif kind == 'H2':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size, bold, italic = 12, True, False
        pf.space_before, pf.space_after = Pt(10), Pt(6)
    elif kind == 'H3':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size, bold, italic = 11, True, False
        pf.space_before, pf.space_after = Pt(8), Pt(4)
    elif kind == 'CAP':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size, bold, italic = 10, True, False
        pf.space_before, pf.space_after = Pt(2), Pt(2)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif kind == 'SRC':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size, bold, italic = 10, False, True
        pf.space_before, pf.space_after = Pt(0), Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        size, bold, italic = 11, False, False
        pf.space_before, pf.space_after = Pt(0), Pt(6)
    if kind not in {'CAP', 'SRC'}:
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in paragraph.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(0, 0, 0)


if not BACKUP_PATH.exists():
    shutil.copy2(DOC_PATH, BACKUP_PATH)
create_architecture_image(ARCH_IMG)
create_erd_image(ERD_IMG)

doc = Document(str(DOC_PATH))
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.5)

find_paragraph(doc, 'Auditoría de seguridad', exact=False).text = PROJECT_TITLE
find_paragraph(doc, 'CUITLÁHUAC, VER.', exact=False).text = COVER_DATE
ag = find_paragraph(doc, 'Apartado opcional destinado a agradecer', exact=False)
ag.text = (TMP / 'agradecimientos.txt').read_text(encoding='utf-8').strip()
res = find_paragraph(doc, 'Breve explicación del contenido del reporte', exact=False)
res.text = (TMP / 'resumen.txt').read_text(encoding='utf-8').strip().split('\n\n')[0]
res2 = res.insert_paragraph_before((TMP / 'resumen.txt').read_text(encoding='utf-8').strip().split('\n\n')[1])
res._element.addnext(res2._element)
find_paragraph(doc, 'CAPITULO 3. DESARROLLO DEL PROYECTO', exact=False).text = 'CAPÍTULO 3. DESARROLLO DEL PROYECTO'
find_paragraph(doc, '2.1 Enfoque y tipo de investigación').style = 'Heading 2'

cap22 = find_paragraph(doc, '2.2 Metodología de desarrollo')
cap3 = find_paragraph(doc, 'CAPÍTULO 3. DESARROLLO DEL PROYECTO')
remove_between(cap22, cap3, include_start=True)
for kind, text in load_items('cap2.txt'):
    p = add_item_before(cap3, kind, text)
    fmt(p, 'H2' if kind == 'H2' else 'H3' if kind == 'H3' else kind)

cap3 = find_paragraph(doc, 'CAPÍTULO 3. DESARROLLO DEL PROYECTO')
cap4 = find_paragraph(doc, 'CAPÍTULO 4. RESULTADOS Y CONCLUSIONES', exact=False)
remove_between(cap3, cap4, include_start=False)
for kind, text in load_items('cap3.txt'):
    p = add_item_before(cap4, kind, text)
    fmt(p, 'H2' if kind == 'H2' else kind)

cap4 = find_paragraph(doc, 'CAPÍTULO 4. RESULTADOS Y CONCLUSIONES', exact=False)
anex = find_paragraph(doc, 'ANEXOS')
remove_between(cap4, anex, include_start=False)
for kind, text in load_items('cap4.txt'):
    p = add_item_before(anex, kind, text)
    fmt(p, 'H2' if kind == 'H2' else kind)

anex = find_paragraph(doc, 'ANEXOS')
bib_sdt = find_next_sdt_after(anex)
body = anex._element.getparent()
children = list(body)
start_idx = children.index(anex._element) + 1
end_idx = children.index(bib_sdt)
for child in children[start_idx:end_idx]:
    body.remove(child)
for p, kind in insert_items_after(anex, load_items('anexos.txt')):
    fmt(p, 'H2' if kind == 'H2' else 'CAP' if kind == 'CAP' else 'SRC' if kind == 'SRC' else kind)

for p in doc.paragraphs:
    txt = p.text.strip()
    style_name = p.style.name if p.style else ''
    if not txt or style_name.startswith('TOC'):
        continue
    if txt in {'AGRADECIMIENTOS','RESUMEN','CAPÍTULO 1. INTRODUCCIÓN','CAPÍTULO 2. METODOLOGÍA','CAPÍTULO 3. DESARROLLO DEL PROYECTO','CAPÍTULO 4. RESULTADOS Y CONCLUSIONES','ANEXOS','BIBLIOGRAFÍA'}:
        p.style = 'Heading 1'
        fmt(p, 'H1')
    elif txt.startswith('Figura '):
        fmt(p, 'CAP')
    elif txt.startswith('Fuente:'):
        fmt(p, 'SRC')
    elif txt.startswith('2.2.1') or txt.startswith('2.2.2') or txt.startswith('2.2.3'):
        p.style = 'Heading 3'
        fmt(p, 'H3')
    elif txt.startswith('2.') or txt.startswith('3.') or txt.startswith('4.') or txt.startswith('Anexo '):
        if not txt.startswith('Figura '):
            p.style = 'Heading 2'
            fmt(p, 'H2')
    elif p._element.getparent().tag.endswith('body') and p._element.getparent() is doc._body._body:
        fmt(p, 'P')

doc.save(str(DOC_PATH))
print('Documento actualizado correctamente')
print(DOC_PATH)

